# Word 文件解析器（docx）

import logging
from io import BytesIO
from .base import FileParser, ParseResult, ContentBlock, ContentType, ExtractionSchema

logger = logging.getLogger(__name__)


class WordParser(FileParser):
    """Word 文件解析器（docx）"""
    
    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}
    
    async def parse(
        self,
        file_bytes: bytes,
        filename: str,
        extraction_schema: ExtractionSchema | None = None,
    ) -> ParseResult:
        """解析 Word 文件"""
        try:
            from docx import Document
            from docx.table import Table
        except ImportError:
            raise ImportError("请安装 python-docx: uv add python-docx")
        
        doc = Document(BytesIO(file_bytes))
        
        blocks = []
        content_parts = []
        tables = []
        
        # 遍历文档元素（保持顺序）
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
            
            if tag == "p":  # 段落
                para = self._find_paragraph(doc, element)
                if para:
                    text = para.text.strip()
                    if text:
                        # 处理标题样式
                        md_text = self._format_paragraph(para, text)
                        content_parts.append(md_text)
                        blocks.append(ContentBlock(
                            type=ContentType.TEXT,
                            content=md_text,
                        ))
            
            elif tag == "tbl":  # 表格
                table = self._find_table(doc, element)
                if table:
                    table_md, table_data = self._parse_table(table)
                    if table_md:
                        content_parts.append(table_md)
                        tables.append(table_data)
                        blocks.append(ContentBlock(
                            type=ContentType.TABLE,
                            content=table_md,
                            raw_data=table_data,
                        ))
        
        return ParseResult(
            content="\n\n".join(content_parts),
            blocks=blocks,
            metadata={
                "format": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "filename": filename,
            },
            tables=tables if tables else None,
        )
    
    def _find_paragraph(self, doc, element):
        """根据 XML 元素查找对应的段落对象"""
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None
    
    def _find_table(self, doc, element):
        """根据 XML 元素查找对应的表格对象"""
        for table in doc.tables:
            if table._element == element:
                return table
        return None
    
    def _format_paragraph(self, para, text: str) -> str:
        """格式化段落（处理标题等样式）"""
        style_name = para.style.name if para.style else ""
        
        # 处理标题
        if style_name.startswith("Heading"):
            # 提取标题级别
            level_str = style_name.replace("Heading", "").strip()
            try:
                level = int(level_str) if level_str else 1
            except ValueError:
                level = 1
            level = min(level, 6)  # 最多 6 级标题
            return f"{'#' * level} {text}"
        
        # 处理列表
        if style_name.startswith("List"):
            return f"- {text}"
        
        return text
    
    def _parse_table(self, table) -> tuple[str, dict]:
        """解析 Word 表格为 Markdown 和结构化数据"""
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                # 合并单元格内的多个段落
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                # 表格内换行替换为 <br>
                cell_text = cell_text.replace("\n", "<br>")
                cells.append(cell_text)
            rows.append(cells)
        
        if not rows:
            return "", {}
        
        # 确定列数（取最大列数）
        max_cols = max(len(row) for row in rows)
        
        # 补齐列数
        for row in rows:
            while len(row) < max_cols:
                row.append("")
        
        # 构建 Markdown 表格
        headers = rows[0]
        md_lines = ["| " + " | ".join(headers) + " |"]
        md_lines.append("|" + " --- |" * max_cols)
        
        for row in rows[1:]:
            md_lines.append("| " + " | ".join(row) + " |")
        
        table_data = {
            "headers": headers,
            "rows": rows[1:],
            "row_count": len(rows) - 1,
        }
        
        return "\n".join(md_lines), table_data
